import docx
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc = docx.Document()
faile = doc.add_paragraph('«Hello World!»','Title')
faile.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


par = doc.add_paragraph('')

par.add_run('\tПо полям, по полям, синий трактор едет к нам')

doc.save('ворд.docx')
print(faile.text,par.text)